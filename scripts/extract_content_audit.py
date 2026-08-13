#!/usr/bin/env python3
"""Extract readable text and basic metadata from campaign source files."""

from __future__ import annotations

import csv
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pdfplumber
from docx import Document
from openpyxl import load_workbook


ROOT = Path("content-audit/01_rozbaleno")
OUTPUT = Path("content-audit/02_extrakce")


def safe_name(path: Path) -> str:
    relative = path.relative_to(ROOT)
    return re.sub(r"[^0-9A-Za-zÀ-ž._-]+", "_", str(relative)).strip("_")


def extract_docx(path: Path) -> str:
    doc = Document(path)
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table_index, table in enumerate(doc.tables, 1):
        chunks.append(f"[TABULKA {table_index}]")
        for row in table.rows:
            chunks.append(" | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells))
    return "\n".join(chunks)


def extract_odt(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("content.xml"))
    texts: list[str] = []
    for element in root.iter():
        tag = element.tag.rsplit("}", 1)[-1]
        if tag in {"p", "h"}:
            text = "".join(element.itertext()).strip()
            if text:
                texts.append(text)
    return "\n".join(texts)


def extract_xlsx(path: Path) -> str:
    workbook = load_workbook(path, data_only=True, read_only=True)
    chunks: list[str] = []
    for sheet in workbook.worksheets:
        chunks.append(f"[LIST: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks)


def extract_pdf(path: Path) -> tuple[str, int]:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for index, page in enumerate(pdf.pages, 1):
            text = (page.extract_text() or "").strip()
            chunks.append(f"[STRANA {index}]\n{text}")
    return "\n\n".join(chunks), page_count


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    inventory: list[dict[str, object]] = []
    for path in sorted(ROOT.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        text = ""
        pages = None
        status = "metadata-only"
        try:
            if suffix == ".docx":
                text = extract_docx(path)
                status = "extracted"
            elif suffix == ".odt":
                text = extract_odt(path)
                status = "extracted"
            elif suffix == ".xlsx":
                text = extract_xlsx(path)
                status = "extracted"
            elif suffix == ".pdf":
                text, pages = extract_pdf(path)
                status = "extracted" if text.strip() else "needs-ocr"
        except Exception as error:  # Preserve a complete audit even if one file is malformed.
            status = f"error: {type(error).__name__}: {error}"

        output_path = None
        if text.strip():
            output_path = OUTPUT / f"{safe_name(path)}.txt"
            output_path.write_text(text, encoding="utf-8")

        inventory.append(
            {
                "source": str(path),
                "type": suffix.lstrip("."),
                "bytes": path.stat().st_size,
                "pages": pages,
                "characters": len(text),
                "words": len(text.split()),
                "status": status,
                "extracted_text": str(output_path) if output_path else "",
            }
        )

    with (OUTPUT / "inventory.json").open("w", encoding="utf-8") as stream:
        json.dump(inventory, stream, ensure_ascii=False, indent=2)
    with (OUTPUT / "inventory.csv").open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=inventory[0].keys())
        writer.writeheader()
        writer.writerows(inventory)

    print(json.dumps({"files": len(inventory), "extracted": sum(bool(row["extracted_text"]) for row in inventory)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
